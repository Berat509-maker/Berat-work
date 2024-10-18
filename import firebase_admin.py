import firebase_admin
from firebase_admin import credentials, db
import tkinter as tk
from tkinter import filedialog, messagebox, ttk, simpledialog
import docx, os

# Path to your Firebase credentials JSON file
cred = credentials.Certificate("./medappmk-93ef2-firebase-adminsdk-qljal-a1f5dd34f4.json")

# Initialize the Firebase app
firebase_admin.initialize_app(cred, {
    'databaseURL': 'https://medappmk-93ef2-default-rtdb.europe-west1.firebasedatabase.app/'  # Replace with your database URL
})

# Initialize the Realtime Database reference
ref = db.reference('/')

def close_app():
    if messagebox.askyesno("Mbyll Aplikacionin", "A doni të mbyllni aplikacionin?"):
        root.destroy()

def new_file():
    if messagebox.askyesno("E Re", "A doni të ruani të dhënat aktuale para se të hapni një skedë të re?"):
        save_data()  # Call save_data to save the current data

    # Clear all entry fields and text boxes
    pacient_id_var.set('')
    emri_var.set('')
    data_lindjes_var.set('')
    grupi_gjakut_var.set('')
    numri_kontrolleve_var.set('')
    data_vizites_var.set('')
    diagnoza_kodi_var.set('')
    diagnoza_pershkrimi_var.set('')
    shenime_text.delete("1.0", tk.END)
    medikament_var.set('')
    doza_var.set('')
    frekuenca_var.set('')

# Function to fetch data from Firebase based on diagnoza_kodi
def fetch_data(diagnoza_kodi):
    users = ref.child('users').order_by_child('diagnoza_kodi').equal_to(diagnoza_kodi).get()
    
    if users:
        for user_id, user_data in users.items():
            # Populate the form with the fetched data
            pacient_id_var.set(user_data.get('pacient_id', ''))
            emri_var.set(user_data.get('emri', ''))
            data_lindjes_var.set(user_data.get('data_lindjes', ''))
            grupi_gjakut_var.set(user_data.get('grupi_gjakut', ''))
            numri_kontrolleve_var.set(user_data.get('numri_kontrolleve', ''))
            data_vizites_var.set(user_data.get('data_vizites', ''))
            diagnoza_pershkrimi_var.set(user_data.get('diagnoza_pershkrimi', ''))
            shenime_text.delete("1.0", tk.END)  # Clear existing text
            shenime_text.insert(tk.END, user_data.get('shenime', ''))
            medikament_var.set(user_data.get('medikament', ''))
            doza_var.set(user_data.get('doza', ''))
            frekuenca_var.set(user_data.get('frekuenca', ''))

            messagebox.showinfo("Sukses", "Të dhënat u morën me sukses!")
            break  # Exit the loop after finding the first match
    else:
        messagebox.showwarning("Këmbanë", "Nuk u gjetën të dhëna për këtë kodi diagnoze.")

# Function to prompt for diagnoza_kodi and fetch data
def gjej_data():
    diagnoza_kodi = simpledialog.askstring("Kodi Diagnozës", "Shkruani Kodi Diagnozës:")
    
    if diagnoza_kodi:
        fetch_data(diagnoza_kodi)

def save_doc():
    global saved_doc_path  # Use global variable to keep track of saved document path

    # Gather the current values from the form
    pacient_id = pacient_id_var.get()
    emri = emri_var.get()
    data_lindjes = data_lindjes_var.get()
    grupi_gjakut = grupi_gjakut_var.get()
    numri_kontrolleve = numri_kontrolleve_var.get()
    data_vizites = data_vizites_var.get()
    diagnoza_kodi = diagnoza_kodi_var.get()
    diagnoza_pershkrimi = diagnoza_pershkrimi_var.get()
    shenime = shenime_text.get("1.0", tk.END).strip()  # Strip any trailing newlines
    shenime_lines = shenime.splitlines()  # Split into individual lines
    shenime_indented = "\n".join(["\t" + line for line in shenime_lines])  # Add tabs to each line
    medikament = medikament_var.get()
    doza = doza_var.get()
    frekuenca = frekuenca_var.get()

    # Create a Word document
    doc = docx.Document()
    
    # Add title and patient info section
    doc.add_heading('Forma Mjekësore', 0)
    doc.add_heading('Informacioni i Pacientit', level=1)

    # Adding information with bold labels and right-aligned values
    def add_left_right_paragraph(label, value):
        # Split the value by newlines to handle multi-line text
        lines = value.split('\n')
        for i, line in enumerate(lines):
            paragraph = doc.add_paragraph()
            paragraph.add_run("\t" * 2)  # Add tabs
            run_label = paragraph.add_run(label)
            run_label.bold = True  # Make the label bold
            paragraph.add_run(line.strip())  # Add the value without leading/trailing whitespace

    add_left_right_paragraph('ID e Pacientit: ', pacient_id)
    add_left_right_paragraph('Emri: ', emri)
    add_left_right_paragraph('Data e Lindjes: ', data_lindjes)
    add_left_right_paragraph('Grupi i Gjakut: ', grupi_gjakut)
    add_left_right_paragraph('Numri i Kontrolleve: ', numri_kontrolleve)
    add_left_right_paragraph('Data e Vizitës: ', data_vizites)

    # Add diagnosis section
    doc.add_heading('Diagnoza', level=1)
    add_left_right_paragraph('Kodi i Diagnozës: ', diagnoza_kodi)
    add_left_right_paragraph('Përshkrimi i Diagnozës: ', diagnoza_pershkrimi)
    add_left_right_paragraph('Shënime: ', shenime_indented)

    # Add treatment section
    doc.add_heading('Trajtimi', level=1)
    add_left_right_paragraph('Medikament: ', medikament)
    add_left_right_paragraph('Doza: ', doza)
    add_left_right_paragraph('Frekuenca: ', frekuenca)

    # Save the document to a file
    file_path = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Word Document", "*.docx"), ("All Files", "*.*")],
        title="Ruaj Dosjen",
        initialfile=f"{pacient_id}-{diagnoza_kodi}-data"  # Default filename
    )
    
    if file_path:
        doc.save(file_path)
        saved_doc_path = file_path  # Store the saved document path
        messagebox.showinfo("Ruaj", f"Dosja u ruajt si {file_path}")
    else:
        messagebox.showwarning("Anullo", "Ruajtja u anulua")



def print_doc():
    if saved_doc_path:
        os.startfile(saved_doc_path, "print")
        messagebox.showinfo("Printo", "Dokumenti po printohet...")
    else:
        messagebox.showwarning("Këmbanë", "Nuk ka dokument për të printuar. Ruani dokumentin së pari.")

# Function to add data to the Realtime Database
def save_data():
    diagnoza_kodi = diagnoza_kodi_var.get()

    # Check for existing diagnosis code
    existing_users = ref.child('users').order_by_child('diagnoza_kodi').equal_to(diagnoza_kodi).get()

    if existing_users:
        messagebox.showwarning("Këmbanë", "Një diagnozë me këtë kod tashmë ekziston!")
        return  # Exit the function to prevent saving

    pacient_id = pacient_id_var.get()
    emri = emri_var.get()
    data_lindjes = data_lindjes_var.get()
    grupi_gjakut = grupi_gjakut_var.get()
    numri_kontrolleve = numri_kontrolleve_var.get()
    data_vizites = data_vizites_var.get()
    diagnoza_pershkrimi = diagnoza_pershkrimi_var.get()
    shenime = shenime_text.get("1.0", tk.END)
    medikament = medikament_var.get()
    doza = doza_var.get()
    frekuenca = frekuenca_var.get()

    # Add data to the Realtime Database
    data = {
        "pacient_id": pacient_id,
        "emri": emri,
        "data_lindjes": data_lindjes,
        "grupi_gjakut": grupi_gjakut,
        "numri_kontrolleve": numri_kontrolleve,
        "data_vizites": data_vizites,
        "diagnoza_kodi": diagnoza_kodi,
        "diagnoza_pershkrimi": diagnoza_pershkrimi,
        "shenime": shenime,
        "medikament": medikament,
        "doza": doza,
        "frekuenca": frekuenca,
    }
    ref.child('users').push(data)  # Push data to 'users' node
    messagebox.showinfo("Sukses!", "Të dhënat u ruajten!")

# Main window setup
root = tk.Tk()
root.title("Forma Mjekësore UI")
root.geometry("1000x750")  # Default window size

# Configure grid weights to allow resizing
root.columnconfigure(0, weight=1)
root.rowconfigure([0, 1, 2, 3], weight=1)

# Applying a modern style to the widgets
style = ttk.Style()
style.configure("TLabel", font=("Arial", 11), padding=5)
style.configure("TButton", font=("Arial", 11), padding=5)
style.configure("TFrame", padding=10)

# Creating a menu bar
menu_bar = tk.Menu(root)

file_menu = tk.Menu(menu_bar, tearoff=0)
file_menu.add_command(label="E Re", command=new_file)  # Connect to new_file
file_menu.add_command(label="Ruaj", command=save_data)
file_menu.add_separator()
file_menu.add_command(label="Printo")
file_menu.add_separator()
file_menu.add_command(label="Dalje", command=close_app)  # Connect to close_app
menu_bar.add_cascade(label="Dosje", menu=file_menu)

# Adding the menu to the window
root.config(menu=menu_bar)

# Variables for storing user input
pacient_id_var = tk.StringVar()
emri_var = tk.StringVar()
data_lindjes_var = tk.StringVar()
grupi_gjakut_var = tk.StringVar()
numri_kontrolleve_var = tk.StringVar()
data_vizites_var = tk.StringVar()
diagnoza_kodi_var = tk.StringVar()
diagnoza_pershkrimi_var = tk.StringVar()
medikament_var = tk.StringVar()
doza_var = tk.StringVar()
frekuenca_var = tk.StringVar()

# Patient info frame (Informacioni i Pacientit)
pacient_frame = ttk.LabelFrame(root, text="Informacioni i Pacientit")
pacient_frame.grid(row=0, column=0, padx=10, pady=10, sticky='nsew')

# Patient labels and entries
ttk.Label(pacient_frame, text="ID e Pacientit:").grid(row=0, column=0, sticky='e')
ttk.Entry(pacient_frame, textvariable=pacient_id_var, width=25).grid(row=0, column=1, padx=10, pady=5, sticky='ew')
ttk.Label(pacient_frame, text="Emri:").grid(row=0, column=2, sticky='e')
ttk.Entry(pacient_frame, textvariable=emri_var, width=25).grid(row=0, column=3, padx=10, pady=5, sticky='ew')
ttk.Label(pacient_frame, text="Data e Lindjes:").grid(row=1, column=0, sticky='e')
ttk.Entry(pacient_frame, textvariable=data_lindjes_var, width=25).grid(row=1, column=1, padx=10, pady=5, sticky='ew')
ttk.Label(pacient_frame, text="Grupi i Gjakut:").grid(row=1, column=2, sticky='e')
ttk.Entry(pacient_frame, textvariable=grupi_gjakut_var, width=25).grid(row=1, column=3, padx=10, pady=5, sticky='ew')
ttk.Label(pacient_frame, text="Numri i Kontrolleve:").grid(row=2, column=0, sticky='e')
ttk.Entry(pacient_frame, textvariable=numri_kontrolleve_var, width=25).grid(row=2, column=1, padx=10, pady=5, sticky='ew')
ttk.Label(pacient_frame, text="Data e Vizitës:").grid(row=2, column=2, sticky='e')
ttk.Entry(pacient_frame, textvariable=data_vizites_var, width=25).grid(row=2, column=3, padx=10, pady=5, sticky='ew')

# Allow the pacient_frame to resize with the window
pacient_frame.columnconfigure([1, 3], weight=1)

# Diagnosis frame (Diagnoza)
diagnoza_frame = ttk.LabelFrame(root, text="Diagnoza")
diagnoza_frame.grid(row=1, column=0, padx=10, pady=10, sticky='nsew')

# Diagnosis labels and entries
ttk.Label(diagnoza_frame, text="Kodi i Diagnozës:").grid(row=0, column=0, sticky='e')
ttk.Entry(diagnoza_frame, textvariable=diagnoza_kodi_var, width=25).grid(row=0, column=1, padx=10, pady=5, sticky='ew')
ttk.Button(diagnoza_frame, text="Merr të Dhënat", command=lambda: fetch_data(diagnoza_kodi_var.get())).grid(row=0, column=2, padx=5, pady=5)

ttk.Label(diagnoza_frame, text="Përshkrimi i Diagnozës:").grid(row=0, column=3, sticky='e')
ttk.Entry(diagnoza_frame, textvariable=diagnoza_pershkrimi_var, width=50).grid(row=0, column=4, padx=10, pady=5, sticky='ew')

# Enlarged Notes Textbox (Shënime)
ttk.Label(diagnoza_frame, text="Shënime:").grid(row=1, column=0, sticky='ne')
shenime_text = tk.Text(diagnoza_frame, height=12, width=60)  # Enlarged text box
shenime_text.grid(row=1, column=1, columnspan=4, padx=10, pady=5, sticky='nsew')

# Make sure diagnosis fields resize properly
diagnoza_frame.columnconfigure([1, 4], weight=1)
diagnoza_frame.rowconfigure(1, weight=1)  # Allow the notes textbox to expand vertically

# Treatment frame (Trajtimi)
trajtimi_frame = ttk.LabelFrame(root, text="Trajtimi")
trajtimi_frame.grid(row=2, column=0, padx=10, pady=10, sticky='nsew')
ttk.Label(trajtimi_frame, text="Medikament:").grid(row=0, column=0, sticky='e')
ttk.Entry(trajtimi_frame, textvariable=medikament_var, width=25).grid(row=0, column=1, padx=10, pady=5, sticky='ew')
ttk.Label(trajtimi_frame, text="Doza:").grid(row=1, column=0, sticky='e')
ttk.Entry(trajtimi_frame, textvariable=doza_var, width=25).grid(row=1, column=1, padx=10, pady=5, sticky='ew')
ttk.Label(trajtimi_frame, text="Frekuenca:").grid(row=2, column=0, sticky='e')
ttk.Entry(trajtimi_frame, textvariable=frekuenca_var, width=25).grid(row=2, column=1, padx=10, pady=5, sticky='ew')

# Allow treatment frame to resize properly
trajtimi_frame.columnconfigure(1, weight=1)

# Action buttons with modern design
button_frame = ttk.Frame(root)
button_frame.grid(row=3, column=0, padx=10, pady=10, sticky='ew')

ttk.Button(button_frame, text="Ruaj në Databazë", command=save_data).grid(row=0, column=0, padx=5, pady=5)
ttk.Button(button_frame, text="Ruaj në Dokument", command=save_doc).grid(row=0, column=1, padx=5, pady=5)
ttk.Button(button_frame, text="Mbyll", command=close_app).grid(row=0, column=2, padx=5, pady=5)


# Make the button frame and window responsive
button_frame.columnconfigure([0, 1], weight=1)
root.columnconfigure(0, weight=1)

# Start the Tkinter event loop
root.mainloop()
